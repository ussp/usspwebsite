#!/usr/bin/env python3
"""Convert the 30-day notice markdown into a Word document.

Reads:  doit-30day-notice-starting-draft-v04222026.md
Writes: doit-30day-notice-starting-draft-v04222026.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

import sys
import glob

_base = sys.argv[1] if len(sys.argv) > 1 else 'doit-30day-notice-detailed-v20260423'
_root = os.path.dirname(__file__)

# If an explicit relative path is passed (e.g., governance/foo), use it; else search subfolders
if os.path.sep in _base or '/' in _base:
    SRC = os.path.join(_root, _base + '.md')
else:
    matches = glob.glob(os.path.join(_root, '**', _base + '.md'), recursive=True)
    if not matches:
        print('ERROR: could not find ' + _base + '.md in ' + _root + ' or its subfolders')
        sys.exit(1)
    SRC = matches[0]

OUT = SRC[:-3] + '.docx'

doc = Document()

# Base style — tight spacing (no extra gaps between paragraphs)
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(4)

# Tighten heading spacing too
for heading_name in ('Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'):
    try:
        h_style = doc.styles[heading_name]
        h_style.paragraph_format.space_before = Pt(8)
        h_style.paragraph_format.space_after = Pt(2)
    except KeyError:
        pass

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
    return h

def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bold_label_then_text(label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    p.add_run(text)
    return p

def parse_inline(text):
    """Return list of (text, bold, italic) tuples from inline markdown (**bold**, *italic*)."""
    # Very small inline parser: handle **bold** and *italic*
    segments = []
    i = 0
    while i < len(text):
        if text.startswith('**', i):
            j = text.find('**', i + 2)
            if j != -1:
                segments.append((text[i+2:j], True, False))
                i = j + 2
                continue
        if text.startswith('*', i):
            j = text.find('*', i + 1)
            if j != -1 and text[i+1:j]:
                segments.append((text[i+1:j], False, True))
                i = j + 1
                continue
        # plain run up to next marker
        next_star = text.find('*', i + 1)
        if next_star == -1:
            segments.append((text[i:], False, False))
            break
        segments.append((text[i:next_star], False, False))
        i = next_star
    return segments

def add_rich_para(text):
    p = doc.add_paragraph()
    for seg_text, bold, italic in parse_inline(text):
        r = p.add_run(seg_text)
        r.bold = bold
        r.italic = italic
    return p

def add_table_from_markdown(rows):
    """rows: list of lists of cells (strings). First row is header."""
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            for seg_text, bold, italic in parse_inline(cell_text):
                run = para.add_run(seg_text)
                run.bold = bold or (r_idx == 0)  # header row is bold
                run.italic = italic
                run.font.size = Pt(10)
    return table

# Parse the markdown file
with open(SRC, 'r', encoding='utf-8') as f:
    lines = f.read().split('\n')

i = 0
skip_block = False  # skip the frontmatter callout at top and internal prep notes
while i < len(lines):
    line = lines[i].rstrip()

    # Skip top frontmatter block (blockquote lines at start)
    if i < 10 and line.startswith('>'):
        i += 1
        continue

    # Skip internal prep notes section entirely (everything after "Internal prep notes")
    if 'Internal prep notes' in line:
        skip_block = True
    if skip_block and 'Email cover for sending this memo' in line:
        skip_block = False
        # Do add the email cover section
    if skip_block:
        i += 1
        continue

    # Horizontal rule
    if re.match(r'^-{3,}$', line):
        i += 1
        continue

    # Table detection (line with pipes, followed by a separator line)
    if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', lines[i+1]):
        # Collect table
        table_rows = []
        while i < len(lines) and '|' in lines[i]:
            l = lines[i].strip()
            if re.match(r'^\|?[\s\-:|]+\|?\s*$', l):
                i += 1
                continue
            # parse cells
            l = l.strip('|').strip()
            cells = [c.strip() for c in re.split(r'\s*\|\s*', l)]
            table_rows.append(cells)
            i += 1
        add_table_from_markdown(table_rows)
        continue

    # Heading
    m = re.match(r'^(#+)\s+(.+)$', line)
    if m:
        level = min(len(m.group(1)), 4)
        text = m.group(2).strip()
        # Strip any leading emoji
        add_heading(text, level=level)
        i += 1
        continue

    # Blockquote
    if line.startswith('> '):
        p = doc.add_paragraph()
        r = p.add_run(line[2:])
        r.italic = True
        p.paragraph_format.left_indent = Inches(0.3)
        i += 1
        continue

    # Bullet
    if line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        for seg_text, bold, italic in parse_inline(line[2:]):
            r = p.add_run(seg_text)
            r.bold = bold
            r.italic = italic
        i += 1
        continue

    # Numbered list
    m = re.match(r'^\d+\.\s+(.+)$', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        for seg_text, bold, italic in parse_inline(m.group(1)):
            r = p.add_run(seg_text)
            r.bold = bold
            r.italic = italic
        i += 1
        continue

    # Code fence — skip
    if line.startswith('```'):
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            i += 1
        i += 1
        continue

    # Blank line — skip (Word paragraphs have natural spacing; no need for blank paragraphs)
    if line.strip() == '':
        i += 1
        continue

    # Regular paragraph
    add_rich_para(line)
    i += 1

doc.save(OUT)
print('Word doc created: ' + OUT)
